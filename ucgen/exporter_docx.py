"""Export UseCaseDocument to Word (.docx) using python-docx."""

from __future__ import annotations

import re
from pathlib import Path

from ucgen import __version__
from ucgen.errors import UCGenError
from ucgen.schema import Entity, UseCaseDocument

# Palette (reference template)
C_NAVY = "1A1A2E"
C_BLUE = "4A90D9"
C_BLUE_DARK = "0F3460"
C_BLUE_MID = "16213E"
C_SUBTITLE = "99CCFF"
C_GREY = "AAAAAA"
C_GREY_TEXT = "666666"
C_ALT_ROW = "E8F0FE"
C_BORDER = "CCCCCC"
C_HEADER_TEXT = "FFFFFF"

CONTENT_WIDTH_IN = 6.5


def _rgb(hex6: str) -> tuple[int, int, int]:
    h = hex6.lstrip("#")
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def _set_cell_shading(cell: object, fill_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc  # noqa: SLF001
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_margins(cell: object, top: int, start: int, bottom: int, end: int) -> None:
    """Set cell margins in twips (dxa)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc  # noqa: SLF001
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for name, val in (
        ("top", top),
        ("left", start),
        ("bottom", bottom),
        ("right", end),
    ):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    tc_pr.append(tc_mar)


def _set_table_full_width(table: object, width_inches: float) -> None:
    from docx.shared import Inches

    table.autofit = False
    table.allow_autofit = False
    table.width = Inches(width_inches)


def _border_cell(cell: object, *, top: str | None = None, bottom: str | None = None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = cell._tc  # noqa: SLF001
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side, color in (("top", top), ("bottom", bottom)):
        if not color:
            continue
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), color)
        tc_borders.append(tag)
    if top or bottom:
        tc_pr.append(tc_borders)


def _font_run(
    run: object,
    name: str,
    size_pt: float,
    *,
    bold: bool = False,
    color_hex: str | None = None,
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color_hex:
        r, g, b = _rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
    r_elm = run._element  # noqa: SLF001
    r_pr = r_elm.rPr
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        r_elm.insert(0, r_pr)
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), name)
        r_fonts.set(qn("w:hAnsi"), name)
        r_pr.append(r_fonts)
    else:
        r_fonts.set(qn("w:ascii"), name)
        r_fonts.set(qn("w:hAnsi"), name)


def _paragraph_border_top(paragraph: object, color_hex: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = paragraph._element  # noqa: SLF001
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "8")
    top.set(qn("w:space"), "2")
    top.set(qn("w:color"), color_hex)
    p_bdr.append(top)
    p_pr.append(p_bdr)


def _add_page_field(paragraph: object) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)  # noqa: SLF001
    run._r.append(instr)  # noqa: SLF001
    run._r.append(fld_char2)  # noqa: SLF001


def _priority_from_markdown(raw: str) -> str:
    m = re.search(r"^priority:\s*(\S+)", raw, re.MULTILINE)
    return m.group(1).strip().strip('"') if m else "medium"


def _heading2(document: object, text: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = True
    r, g, b = _rgb(C_BLUE_DARK)
    run.font.color.rgb = RGBColor(r, g, b)
    p_pr = p._element.get_or_add_pPr()  # noqa: SLF001
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:color"), C_BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _build_title_page(document: object, doc: UseCaseDocument) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    meta = doc.metadata
    priority = _priority_from_markdown(doc.raw_markdown)

    table = document.add_table(rows=1, cols=1)
    _set_table_full_width(table, CONTENT_WIDTH_IN)
    cell = table.rows[0].cells[0]
    _set_cell_shading(cell, C_NAVY)
    _set_cell_margins(cell, 360, 180, 360, 180)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = cp.add_run("ucgen")
    _font_run(r1, "Arial", 36, bold=True, color_hex="FFFFFF")
    cp.add_run().add_break()
    r2 = cp.add_run("Use Case Document")
    _font_run(r2, "Arial", 14, bold=False, color_hex=C_SUBTITLE)

    p_rule = document.add_paragraph()
    p_rule.paragraph_format.space_after = Pt(6)
    p_b = p_rule._element.get_or_add_pPr()  # noqa: SLF001
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p_bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "18")
    bot.set(qn("w:color"), C_BLUE)
    p_bdr.append(bot)
    p_b.append(p_bdr)

    p_id = document.add_paragraph()
    p_id.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_id = p_id.add_run(meta.uc_id)
    _font_run(r_id, "Arial", 26, bold=True, color_hex=C_BLUE)

    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_title.add_run(meta.title)
    _font_run(r_t, "Arial", 20, bold=True, color_hex=C_NAVY)

    p_thin = document.add_paragraph()
    p_thin.paragraph_format.space_after = Pt(6)
    pb = p_thin._element.get_or_add_pPr()  # noqa: SLF001
    pbd = OxmlElement("w:pBdr")
    b2 = OxmlElement("w:bottom")
    b2.set(qn("w:val"), "single")
    b2.set(qn("w:sz"), "4")
    b2.set(qn("w:color"), C_BORDER)
    pbd.append(b2)
    pb.append(pbd)

    p_meta = document.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = p_meta.add_run(f"{meta.actor}  ·  {meta.domain}  ·  {priority}  ·  {meta.goal_level}")
    _font_run(rm, "Arial", 11, color_hex=C_GREY_TEXT)

    ts = doc.generated_at.isoformat()
    p_ts = document.add_paragraph()
    p_ts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ts = p_ts.add_run(f"{ts}  ·  {doc.provider} / {doc.model}")
    _font_run(r_ts, "Arial", 8, color_hex=C_GREY)

    document.add_page_break()


def _set_body_header_footer(section: object, uc_id: str, title: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    header = section.header
    ht = header.add_table(rows=1, cols=2, width=Inches(CONTENT_WIDTH_IN))
    _set_table_full_width(ht, CONTENT_WIDTH_IN)
    ht.columns[0].width = Inches(4.5)
    ht.columns[1].width = Inches(2.0)
    c0 = ht.rows[0].cells[0]
    c1 = ht.rows[0].cells[1]
    _set_cell_shading(c0, C_NAVY)
    _set_cell_shading(c1, C_BLUE)
    _set_cell_margins(c0, 80, 115, 80, 115)
    _set_cell_margins(c1, 80, 115, 80, 115)
    p0 = c0.paragraphs[0]
    r0 = p0.add_run("ucgen  ·  Use Case Document")
    _font_run(r0, "Arial", 9, bold=True, color_hex="FFFFFF")
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p1.add_run(f"v{__version__}")
    _font_run(r1, "Arial", 8, bold=True, color_hex="FFFFFF")

    footer = section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    _paragraph_border_top(fp, C_BLUE)
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lead = fp.add_run(f"{uc_id} — {title}   |   Page ")
    _font_run(lead, "Arial", 8, color_hex=C_GREY)
    _add_page_field(fp)


def _section_overview(document: object, doc: UseCaseDocument) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    meta = doc.metadata
    priority = _priority_from_markdown(doc.raw_markdown)

    _heading2(document, "1 — Overview")

    t = document.add_table(rows=1, cols=3)
    _set_table_full_width(t, CONTENT_WIDTH_IN)
    for i, (text, color) in enumerate(
        (
            (meta.uc_id, C_NAVY),
            (f"Priority: {priority}", C_BLUE_MID),
            (f"Goal Level: {meta.goal_level}", C_BLUE_DARK),
        )
    ):
        cell = t.rows[0].cells[i]
        _set_cell_shading(cell, color)
        _set_cell_margins(cell, 100, 115, 100, 115)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        _font_run(r, "Arial", 10, bold=True, color_hex=C_HEADER_TEXT)

    rows_data = [
        ("Use Case Name", meta.title),
        ("Actor", meta.actor),
        ("Supporting Actors", ", ".join(meta.supporting_actors) or "—"),
        ("Domain", meta.domain),
        ("System Boundary", meta.system_boundary or "—"),
        ("Trigger", meta.trigger),
        ("Goal", meta.goal),
    ]
    info = document.add_table(rows=len(rows_data), cols=2)
    _set_table_full_width(info, CONTENT_WIDTH_IN)
    info.columns[0].width = Inches(1.8)
    info.columns[1].width = Inches(4.7)
    for idx, (label, value) in enumerate(rows_data):
        c_a = info.rows[idx].cells[0]
        c_b = info.rows[idx].cells[1]
        fill = C_ALT_ROW if idx % 2 else "FFFFFF"
        _set_cell_shading(c_a, fill)
        _set_cell_shading(c_b, fill)
        _set_cell_margins(c_a, 72, 115, 72, 115)
        _set_cell_margins(c_b, 72, 115, 72, 115)
        _border_cell(c_a, bottom=C_BORDER)
        _border_cell(c_b, bottom=C_BORDER)
        p_a = c_a.paragraphs[0]
        r_a = p_a.add_run(label)
        _font_run(r_a, "Arial", 10, bold=True)
        p_b = c_b.paragraphs[0]
        r_b = p_b.add_run(value)
        _font_run(r_b, "Arial", 10)


def _bullet_list(document: object, items: list[str], *, bullet_color: str) -> None:
    from docx.shared import Inches, Pt

    for item in items:
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(4)
        r_b = p.add_run("▸  ")
        _font_run(r_b, "Arial", 10, color_hex=bullet_color)
        r_t = p.add_run(item)
        _font_run(r_t, "Arial", 10, color_hex=C_GREY_TEXT)


def _section_conditions(document: object, doc: UseCaseDocument) -> None:

    sec = doc.sections
    _heading2(document, "2 — Conditions")
    p_pre = document.add_paragraph()
    r_pre = p_pre.add_run("Preconditions")
    _font_run(r_pre, "Arial", 11, bold=True, color_hex=C_BLUE_DARK)
    _bullet_list(document, sec.preconditions, bullet_color=C_BLUE)
    p_post = document.add_paragraph()
    r_post = p_post.add_run("Postconditions")
    _font_run(r_post, "Arial", 11, bold=True, color_hex=C_BLUE_DARK)
    _bullet_list(document, sec.postconditions, bullet_color=C_BLUE)

    p_m = document.add_paragraph()
    rm = p_m.add_run("Minimal Guarantee")
    _font_run(rm, "Arial", 11, bold=True, color_hex=C_BLUE_DARK)
    pm = document.add_paragraph()
    r_min = pm.add_run(sec.minimal_guarantee)
    _font_run(r_min, "Arial", 10)

    p_s = document.add_paragraph()
    rs = p_s.add_run("Success Guarantee")
    _font_run(rs, "Arial", 11, bold=True, color_hex=C_BLUE_DARK)
    ps = document.add_paragraph()
    r_succ = ps.add_run(sec.success_guarantee)
    _font_run(r_succ, "Arial", 10)


def _section_normal_course(document: object, doc: UseCaseDocument) -> None:
    from docx.shared import Inches

    _heading2(document, "3 — Normal Course")
    steps = doc.sections.normal_course
    if not steps:
        p = document.add_paragraph()
        _font_run(p.add_run("—"), "Arial", 10)
        return
    t = document.add_table(rows=1 + len(steps), cols=4)
    _set_table_full_width(t, CONTENT_WIDTH_IN)
    widths = (Inches(0.6), Inches(1.1), Inches(2.4), Inches(2.4))
    for i, w in enumerate(widths):
        t.columns[i].width = w
    headers = ("Step", "Actor", "Action", "System Response")
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        _font_run(r, "Arial", 10, bold=True, color_hex=C_HEADER_TEXT)
        _set_cell_shading(cell, C_BLUE_DARK)
        _set_cell_margins(cell, 80, 115, 80, 115)
    for ri, step in enumerate(steps, start=1):
        fill = "FFFFFF" if ri % 2 else C_ALT_ROW
        vals = (str(step.step), step.actor, step.action, step.system_response)
        for ci, val in enumerate(vals):
            cell = t.rows[ri].cells[ci]
            _set_cell_shading(cell, fill)
            _set_cell_margins(cell, 72, 115, 72, 115)
            _border_cell(cell, bottom=C_BORDER)
            p = cell.paragraphs[0]
            _font_run(p.add_run(val), "Arial", 10)


def _section_alternative(document: object, doc: UseCaseDocument) -> None:
    from docx.shared import Inches

    _heading2(document, "4 — Alternative Courses")
    alts = doc.sections.alternative_courses
    if not alts:
        p = document.add_paragraph()
        _font_run(p.add_run("—"), "Arial", 10)
        return
    t = document.add_table(rows=len(alts), cols=2)
    _set_table_full_width(t, CONTENT_WIDTH_IN)
    t.columns[0].width = Inches(3.25)
    t.columns[1].width = Inches(3.25)
    for idx, alt in enumerate(alts):
        left = f"Ref {alt.ref}: {alt.condition}"
        for ci, val in enumerate((left, alt.response)):
            cell = t.rows[idx].cells[ci]
            fill = "FFFFFF" if idx % 2 else C_ALT_ROW
            _set_cell_shading(cell, fill)
            _set_cell_margins(cell, 72, 115, 72, 115)
            _border_cell(cell, bottom=C_BORDER)
            p = cell.paragraphs[0]
            _font_run(p.add_run(val), "Arial", 10)


def _section_nfr(document: object, doc: UseCaseDocument) -> None:
    from docx.shared import Inches

    _heading2(document, "5 — Non-Functional Requirements")
    rows = doc.sections.nfr or []
    if not rows:
        p = document.add_paragraph()
        _font_run(p.add_run("—"), "Arial", 10)
        return
    t = document.add_table(rows=1 + len(rows), cols=3)
    _set_table_full_width(t, CONTENT_WIDTH_IN)
    t.columns[0].width = Inches(1.3)
    t.columns[1].width = Inches(3.5)
    t.columns[2].width = Inches(1.7)
    hdr = ("Type", "Requirement", "Threshold")
    for i, h in enumerate(hdr):
        cell = t.rows[0].cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        _font_run(r, "Arial", 10, bold=True, color_hex=C_HEADER_TEXT)
        _set_cell_shading(cell, C_BLUE_DARK)
        _set_cell_margins(cell, 80, 115, 80, 115)
    for ri, n in enumerate(rows, start=1):
        vals = (n.type, n.requirement, n.threshold or "—")
        for ci, val in enumerate(vals):
            cell = t.rows[ri].cells[ci]
            fill = "FFFFFF" if ri % 2 else C_ALT_ROW
            _set_cell_shading(cell, fill)
            _set_cell_margins(cell, 72, 115, 72, 115)
            _border_cell(cell, bottom=C_BORDER)
            p = cell.paragraphs[0]
            run = p.add_run(val)
            if ci == 2:
                _font_run(run, "Arial", 10, color_hex="10B981")
            else:
                _font_run(run, "Arial", 10)


def _section_info_req(document: object, doc: UseCaseDocument) -> None:
    from docx.shared import Inches

    _heading2(document, "6 — Information Requirements")
    reqs = doc.sections.information_requirements
    if not reqs:
        p = document.add_paragraph()
        _font_run(p.add_run("—"), "Arial", 10)
        return
    t = document.add_table(rows=1 + len(reqs), cols=3)
    _set_table_full_width(t, CONTENT_WIDTH_IN)
    t.columns[0].width = Inches(0.8)
    t.columns[1].width = Inches(3.7)
    t.columns[2].width = Inches(2.0)
    hdr = ("Step", "Data Needed", "Format")
    for i, h in enumerate(hdr):
        cell = t.rows[0].cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        _font_run(r, "Arial", 10, bold=True, color_hex=C_HEADER_TEXT)
        _set_cell_shading(cell, C_BLUE_DARK)
        _set_cell_margins(cell, 80, 115, 80, 115)
    for ri, req in enumerate(reqs, start=1):
        fmt = req.format or req.source or "—"
        vals = (str(req.step), req.data_needed, fmt)
        for ci, val in enumerate(vals):
            cell = t.rows[ri].cells[ci]
            fill = "FFFFFF" if ri % 2 else C_ALT_ROW
            _set_cell_shading(cell, fill)
            _set_cell_margins(cell, 72, 115, 72, 115)
            _border_cell(cell, bottom=C_BORDER)
            p = cell.paragraphs[0]
            _font_run(p.add_run(str(val)), "Arial", 10)


def _entity_fields_text(entity: Entity) -> str:
    parts: list[str] = []
    for f in entity.fields:
        line = f"{f.name}:{f.type}"
        if f.constraints:
            line += f" [{', '.join(f.constraints)}]"
        parts.append(line)
    return ", ".join(parts) if parts else "—"


def _section_entities(document: object, doc: UseCaseDocument) -> None:
    from docx.shared import Inches

    _heading2(document, "7 — Domain Entities")
    ents = doc.entities.entities
    if not ents:
        p = document.add_paragraph()
        _font_run(p.add_run("—"), "Arial", 10)
        return
    t = document.add_table(rows=1 + len(ents), cols=3)
    _set_table_full_width(t, CONTENT_WIDTH_IN)
    t.columns[0].width = Inches(1.5)
    t.columns[1].width = Inches(3.0)
    t.columns[2].width = Inches(2.0)
    hdr = ("Entity", "Fields (name:type)", "Relationships")
    for i, h in enumerate(hdr):
        cell = t.rows[0].cells[i]
        p = cell.paragraphs[0]
        r = p.add_run(h)
        _font_run(r, "Arial", 10, bold=True, color_hex=C_HEADER_TEXT)
        _set_cell_shading(cell, C_NAVY)
        _set_cell_margins(cell, 80, 115, 80, 115)
    for ri, ent in enumerate(ents, start=1):
        rel = ", ".join(ent.relationships) if ent.relationships else "—"
        vals = (ent.name, _entity_fields_text(ent), rel)
        for ci, val in enumerate(vals):
            cell = t.rows[ri].cells[ci]
            fill = "FFFFFF" if ri % 2 else C_ALT_ROW
            _set_cell_shading(cell, fill)
            _set_cell_margins(cell, 72, 115, 72, 115)
            _border_cell(cell, bottom=C_BORDER)
            p = cell.paragraphs[0]
            _font_run(p.add_run(val), "Arial", 10)


def _section_open_issues(document: object, doc: UseCaseDocument) -> None:

    _heading2(document, "8 — Open Issues")
    issues = doc.sections.open_issues or []
    if not issues:
        p = document.add_paragraph()
        _font_run(p.add_run("—"), "Arial", 10)
        return
    _bullet_list(document, issues, bullet_color=C_BLUE)


def _metadata_footer(document: object, doc: UseCaseDocument) -> None:
    from docx.shared import Pt

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    txt = (
        f"Generated by ucgen v{__version__}  ·  Provider: {doc.provider}  ·  "
        f"Duration: {doc.duration_ms}ms"
    )
    _font_run(p.add_run(txt), "Arial", 8, color_hex=C_GREY)


def export_docx(document: UseCaseDocument, output_path: Path) -> Path:
    """Build a Word document from a use case and write it to ``output_path``."""
    try:
        from docx import Document
        from docx.enum.section import WD_SECTION
        from docx.shared import Inches
    except ImportError as exc:
        raise ImportError("python-docx not installed. Run: pip install python-docx") from exc

    output_path = Path(output_path)
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)

        docx_doc = Document()
        section0 = docx_doc.sections[0]
        section0.page_height = Inches(11)
        section0.page_width = Inches(8.5)
        section0.left_margin = Inches(1)
        section0.right_margin = Inches(1)
        section0.top_margin = Inches(1)
        section0.bottom_margin = Inches(1)

        _build_title_page(docx_doc, document)

        body_section = docx_doc.add_section(WD_SECTION.NEW_PAGE)
        body_section.page_height = Inches(11)
        body_section.page_width = Inches(8.5)
        body_section.left_margin = Inches(1)
        body_section.right_margin = Inches(1)
        body_section.top_margin = Inches(1)
        body_section.bottom_margin = Inches(1)
        _set_body_header_footer(body_section, document.metadata.uc_id, document.metadata.title)

        _section_overview(docx_doc, document)
        _section_conditions(docx_doc, document)
        _section_normal_course(docx_doc, document)
        _section_alternative(docx_doc, document)
        _section_nfr(docx_doc, document)
        _section_info_req(docx_doc, document)
        _section_entities(docx_doc, document)
        _section_open_issues(docx_doc, document)
        _metadata_footer(docx_doc, document)

        docx_doc.save(str(output_path))
        return output_path.resolve()
    except Exception as exc:
        raise UCGenError(message=f"Word document export failed: {exc}") from exc
